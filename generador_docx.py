"""
Módulo para generar el informe en formato DOCX
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
from textos import (
    PARRAFOS_FIJOS, PARRAFO_A, PARRAFO_C, PARRAFO_O, PARRAFO_F, PARRAFO_TR, PARRAFO_ejecutivo, PARRAFO_alerta, PARRAFO_orientacion,
)

def agregar_portada(doc: Document, nombre_completo: str, datos: dict) -> None:
    """
    Agrega la portada del informe
    
    Args:
        doc: Documento de Word
        nombre_completo: Nombre completo del encuestado
        datos: Diccionario con datos generales (edad, fecha_aplicacion, etc.)
    """
    # Título
    titulo = doc.add_heading('Prueba de Redes Atencionales ANT', 0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Aumentar tamaño de fuente del título
    for run in titulo.runs:
        run.font.size = Pt(24)
    
    # Espacio
    doc.add_paragraph().add_run().add_break()
    
    # Información del participante
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Nombre del encuestado
    nombre_run = info.add_run(f"Nombre del encuestado: {nombre_completo}\n")
    nombre_run.bold = True
    nombre_run.font.size = Pt(14)
    
    # Edad
    if datos.get('edad'):
        edad_run = info.add_run(f"Edad: {datos['edad']} años\n")
        edad_run.font.size = Pt(14)
    
    # Fecha de aplicación (si está disponible)
    if datos.get('fecha_aplicacion'):
        fecha_app_run = info.add_run(f"Fecha de aplicación: {datos['fecha_aplicacion']}\n")
        fecha_app_run.font.size = Pt(14)
    
    # Fecha del informe
    fecha_actual = datetime.now().strftime("%d/%m/%Y")
    fecha_informe_run = info.add_run(f"Fecha del informe: {fecha_actual}\n")
    fecha_informe_run.font.size = Pt(14)
    
    # Espacio
    doc.add_paragraph().add_run().add_break()
    
    # Nota confidencial
    nota = doc.add_paragraph()
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nota_run = nota.add_run(f"Este es un informe de evaluación cognitiva, obtenido a partir del rendimiento de {nombre_completo} en la prueba ANT (Test de Redes Atencionales).")
    doc.add_paragraph()  # Espacio
    nota.add_run("\n" + "-" * 50 + "\n")  # Línea de separación
    nota.add_run("\nEste es un informe confidencial de carácter educativo u orientativo. No es un diagnóstico clínico y su interpretación es conveniente la realice un profesional competente.")
    nota_run.italic = True
    nota_run.font.size = Pt(12)

def crear_informe_docx(resultados, clasificaciones, nombre_caso="caso", scale_factor_width=0.8, scale_factor_height=0.5):
    """
    Crea el documento Word con el informe del test ANT
    
    Args:
        resultados: Dict con puntuaciones directas y datos del encuestado
        clasificaciones: Dict con clasificaciones (PT)
        nombre_caso: Nombre del evaluado (fallback si no está en resultados)
        
    Returns:
        Document: Documento Word generado
    """
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Extraer nombres de resultados
    nombre_completo = resultados.get('nombre_completo') or nombre_caso
    nombre = resultados.get('nombre') or nombre_caso
    
    # Preparar datos para la portada
    datos_portada = {
        'edad': resultados.get('edad'),
        'fecha_aplicacion': resultados.get('fecha_aplicacion')
    }
    
    # 1. PORTADA (Primera página)
    agregar_portada(doc, nombre_completo, datos_portada)
    doc.add_page_break()
    
    # ========================================================================
    # TÍTULO E INTRODUCCIÓN
    # ========================================================================
    # ========================================================================
    # DESCRIPCIÓN DE LA PRUEBA
    # ========================================================================
    
    parrafo = doc.add_paragraph()
    run = parrafo.add_run(PARRAFOS_FIJOS['titulo_general_prueba'])
    run.bold = True
    doc.add_paragraph(PARRAFOS_FIJOS['objetivo_prueba'])

    parrafo = doc.add_paragraph()
    run = parrafo.add_run(PARRAFOS_FIJOS['titulo_procedimiento'])
    run.bold = True
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_procedimiento1'])
    
    # Imagen parte spatial cueing
    script_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'images')
    grafico_path = os.path.join(script_dir, 'spatial cue.png')
    if os.path.exists(grafico_path):
        try:
            # Obtener dimensiones de la imagen
            img = Image.open(grafico_path)
            width, height = img.size
            # Aplicar factor de escala
            new_width = Inches(width / 120 * scale_factor_width)  # 96 DPI
            new_height = Inches(height / 110 * scale_factor_height)
            # Agregar imagen
            doc.add_picture(grafico_path, width=new_width, height=new_height)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            print(f"Advertencia: No se pudo insertar la imagen spatial cue.png: {e}")
    
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_procedimiento2'])
        
    # Imagen parte PC
    grafico_path = os.path.join(script_dir, 'congruency cue.png')
    if os.path.exists(grafico_path):
        try:
            img = Image.open(grafico_path)
            width, height = img.size
            new_width = Inches(width / 120 * scale_factor_width)  # 96 DPI
            new_height = Inches(height / 110 * scale_factor_height)
            doc.add_picture(grafico_path, width=new_width, height=new_height)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            print(f"Advertencia: No se pudo insertar la imagen congruency cue.png: {e}")

    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_procedimiento3'])

    parrafo = doc.add_paragraph()
    run = parrafo.add_run(PARRAFOS_FIJOS['titulo_indices'])
    run.bold = True
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_indices'])

    # ========================================================================
    # INSERTAR Tabla resultados
    # ========================================================================
    
    # Título de resultados
    titulo_resultados = doc.add_paragraph()
    titulo_resultados.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo_resultados.add_run(PARRAFOS_FIJOS['titulo_resultados'].format(nombre_completo=nombre_completo))
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph()  # Espacio
    
    doc.add_paragraph(PARRAFOS_FIJOS['texto_tabla_resultados'].format(nombre=nombre, nombre_completo=nombre_completo))

    # Tabla PDs, PTs y clasificaciones
    tabla = doc.add_table(rows=4, cols=10)

    # Aplicar estilo simple con bordes negros y encabezado con fondo gris claro
    tabla.style = 'Table Grid'
    # Aplicar fondo gris claro al encabezado
    for cell in tabla.rows[0].cells:
        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
    
    # Encabezados
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Aciertos'
    hdr_cells[2].text = 'Comisiones'
    hdr_cells[3].text = 'Omisiones'
    hdr_cells[4].text = 'Fatiga precisión'
    hdr_cells[5].text = 'Fatiga velocidad'
    hdr_cells[6].text = 'Velocidad'
    hdr_cells[7].text = 'Alerta'
    hdr_cells[8].text = 'Orientación'
    hdr_cells[9].text = 'Control ejecutivo'


    # Centrar el texto en las celdas del encabezado
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    # Fila 1: PD (Puntuación Directa)
    row1 = tabla.rows[1].cells
    row1[0].text = 'Punt. Directa'
    row1[1].text = str(round(resultados.get('PD_A', 0), 4))
    row1[2].text = str(round(resultados.get('PD_C', 0), 4))
    row1[3].text = str(round(resultados.get('PD_O', 0), 4))
    row1[4].text = str(round(resultados.get('PD_F_A', 0), 4))
    row1[5].text = str(round(resultados.get('PD_F_TR', 0), 4))
    row1[6].text = str(round(resultados.get('PD_TR', 0), 4))
    row1[7].text = str(round(resultados.get('PD_TR_alerta', 0), 4))
    row1[8].text = str(round(resultados.get('PD_TR_orientacion', 0), 4))
    row1[9].text = str(round(resultados.get('PD_TR_ejecutivo', 0), 4))
    
    # Fila 2: PT (Puntuación Típica)
    row2 = tabla.rows[2].cells
    row2[0].text = 'Punt. Típica'
    row2[1].text = str(resultados.get('PT_A', 0))
    row2[2].text = str(resultados.get('PT_C', 0))
    row2[3].text = str(resultados.get('PT_O', 0))
    row2[4].text = str(resultados.get('PT_F_A', 0))
    row2[5].text = str(resultados.get('PT_TR', 0))
    row2[6].text = str(resultados.get('PT_F_TR', 0))
    row2[7].text = str(resultados.get('PT_TR_alerta', 0))
    row2[8].text = str(resultados.get('PT_TR_orientacion', 0))
    row2[9].text = str(resultados.get('PT_TR_ejecutivo', 0))
    
    # Fila 3: Clasificación (bajo/normal/alto)
    row3 = tabla.rows[3].cells
    row3[0].text = 'Rendimiento'
    row3[1].text = str(resultados.get('Clasificacion_A', '-'))
    row3[2].text = str(resultados.get('Clasificacion_C', '-'))
    row3[3].text = str(resultados.get('Clasificacion_O', '-'))
    row3[4].text = str(resultados.get('Clasificacion_F_A', '-'))
    row3[5].text = str(resultados.get('Clasificacion_F_TR', '-'))
    row3[6].text = str(resultados.get('Clasificacion_TR', '-'))
    row3[7].text = str(resultados.get('Clasificacion_TR_alerta', '-'))
    row3[8].text = str(resultados.get('Clasificacion_TR_orientacion', '-'))
    row3[9].text = str(resultados.get('Clasificacion_TR_ejecutivo', '-'))
    
    # Colorear celdas según el rendimiento
    # Verde si es alto, rojo si es bajo (columnas: Aciertos, Alerta, Orientación, Ejecutivo)
    green_if_high = [1, 4, 5, 9]
    # Rojo si es alto, verde si es bajo (columnas: Comisiones, Omisiones, Fatiga precisión, Fatiga velocidad, Velocidad)
    red_if_high = [2, 3, 4, 5, 6, 7, 8]
    
    for idx in green_if_high:
        clasificacion = str(row3[idx].text).lower()
        if clasificacion == 'alto':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="90EE90"/>'.format(nsdecls('w'))))
        elif clasificacion == 'bajo':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF6B6B"/>'.format(nsdecls('w'))))
    
    for idx in red_if_high:
        clasificacion = str(row3[idx].text).lower()
        if clasificacion == 'alto':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF6B6B"/>'.format(nsdecls('w'))))
        elif clasificacion == 'bajo':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="90EE90"/>'.format(nsdecls('w'))))

    # Centrar el texto en las celdas de las filas 1, 2 y 3
    for row in [row1, row2, row3]:
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Espacio

    # ========================================================================
    # PÁRRAFOS CONDICIONALES
    # ========================================================================

    # Construir clave para buscar el párrafo condicional A, C, O 
    clas_A = resultados.get('Clasificacion_A', 'normal')
    clas_C = resultados.get('Clasificacion_C', 'normal')
    clas_O = resultados.get('Clasificacion_O', 'normal')
    clas_F_A = resultados.get('Clasificacion_F_A', 'normal')
    clas_F_TR = resultados.get('Clasificacion_F_TR', 'normal')
    clas_TR = resultados.get('Clasificacion_TR', 'normal')
    clas_TR_alerta = resultados.get('Clasificacion_TR_alerta', 'normal')
    clas_TR_orientacion = resultados.get('Clasificacion_TR_orientacion', 'normal')
    clas_TR_ejecutivo = resultados.get('Clasificacion_TR_ejecutivo', 'normal')
    

 # Construir el texto interpretativo usando los diccionarios importados
    texto_condicional = ""

   # Añadir párrafo para A (Aciertos)
    clave_A = f'A {clas_A}'
    if clave_A in PARRAFO_A:
        texto_condicional += PARRAFO_A[clave_A].format(nombre=nombre) + " "
    
    # Añadir párrafo para C (Comisiones)
    clave_C = f'C {clas_C}'
    if clave_C in PARRAFO_C:
        texto_condicional += PARRAFO_C[clave_C].format(nombre=nombre) + " "
    
    # Añadir párrafo para O (Omisiones)
    clave_O = f'O {clas_O}'
    if clave_O in PARRAFO_O:
        texto_condicional += PARRAFO_O[clave_O].format(nombre=nombre) + " "
    
    # Añadir párrafo para F (Fatiga) - combinar F_A y F_TR
    clave_F = f'F_A {clas_F_A} y F_TR {clas_F_TR}'
    if clave_F in PARRAFO_F:
        texto_condicional += PARRAFO_F[clave_F].format(nombre=nombre) + " "
    
    # Salto de línea entre párrafos
    texto_condicional += "\n"
    
    # Añadir párrafo para TR (Tiempo de Reacción)
    clave_TR = f'TR {clas_TR}'
    if clave_TR in PARRAFO_TR:
        texto_condicional += PARRAFO_TR[clave_TR].format(nombre=nombre) + " "
    
    # Añadir párrafo para TR_alerta (Red de Alerta)
    clave_alerta = f'TR_alerta {clas_TR_alerta}'
    if clave_alerta in PARRAFO_alerta:
        texto_condicional += PARRAFO_alerta[clave_alerta].format(nombre=nombre) + " "
    
    # Añadir párrafo para TR_orientacion (Red de Orientación)
    clave_orientacion = f'TR_orientacion {clas_TR_orientacion}'
    if clave_orientacion in PARRAFO_orientacion:
        texto_condicional += PARRAFO_orientacion[clave_orientacion].format(nombre=nombre) + " "
    
    # Añadir párrafo para TR_ejecutivo (Red Ejecutiva)
    clave_ejecutivo = f'TR_ejecutivo {clas_TR_ejecutivo}'
    if clave_ejecutivo in PARRAFO_ejecutivo:
        texto_condicional += PARRAFO_ejecutivo[clave_ejecutivo].format(nombre=nombre) + " "

  # Añadir todo el texto condicional al documento en un solo párrafo
    if texto_condicional:
        doc.add_paragraph(texto_condicional)
    else:
        doc.add_paragraph(f"Resultados del análisis de {nombre}: A={clas_A}, C={clas_C}, O={clas_O}, F_A={clas_F_A}, F_TR={clas_F_TR}, TR={clas_TR}.")


    # ========================================================================
    return doc

def guardar_informe(doc, ruta_salida):
    """
    Guarda el documento generado
    
    Args:
        doc: Documento Word
        ruta_salida: Ruta donde guardar el archivo
    """
    doc.save(ruta_salida)
    print(f"Informe generado exitosamente en: {ruta_salida}")
